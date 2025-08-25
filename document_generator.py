#!/usr/bin/env python3
"""
Document Generator for TalentCraft Marketing Content
Creates professionally formatted Word documents that mirror Template1.pdf structure
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime

class TalentCraftDocumentGenerator:
    """Generate professionally formatted Word documents matching Template1.pdf style"""
    
    def __init__(self):
        self.doc = None
        self._setup_styles()
    
    def _setup_styles(self):
        """Initialize document with custom styles matching Template1.pdf"""
        self.doc = Document()
        
        # Set up document margins - optimized for single page
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)  # Reduced for single page fit
            section.bottom_margin = Inches(0.75)  # Reduced for single page fit
            section.left_margin = Inches(0.85)  # Slightly reduced
            section.right_margin = Inches(0.85)  # Slightly reduced
            
            # Set page orientation and size
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
        
        # Create custom styles matching Template1.pdf aesthetic
        styles = self.doc.styles
        
        # Header logo style for top section
        header_style = styles.add_style('Custom Header', WD_STYLE_TYPE.PARAGRAPH)
        header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_style.paragraph_format.space_after = Pt(18)
        header_style.paragraph_format.space_before = Pt(6)
        
        # Main title style - impactful but optimized for single page
        title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Calibri'  # Clean, professional font
        title_font.size = Pt(20)  # Slightly reduced for single page
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 0, 0)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(8)  # Reduced for single page
        title_style.paragraph_format.space_before = Pt(12)  # Reduced for single page
        title_style.paragraph_format.line_spacing = 1.15
        title_style.paragraph_format.keep_with_next = True
        
        # Subtitle style - professional but single-page optimized
        subtitle_style = styles.add_style('Custom Subtitle', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_font = subtitle_style.font
        subtitle_font.name = 'Calibri'
        subtitle_font.size = Pt(13)  # Slightly reduced for single page
        subtitle_font.bold = False
        subtitle_font.color.rgb = RGBColor(64, 64, 64)  # Darker gray for better contrast
        subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_style.paragraph_format.space_after = Pt(10)  # Reduced for single page
        subtitle_style.paragraph_format.space_before = Pt(4)  
        subtitle_style.paragraph_format.line_spacing = 1.15
        
        # Divider line style - minimal for single page
        divider_style = styles.add_style('Custom Divider', WD_STYLE_TYPE.PARAGRAPH)
        divider_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        divider_style.paragraph_format.space_before = Pt(2)  # Minimal spacing
        divider_style.paragraph_format.space_after = Pt(4)  # Minimal spacing
        
        # Introduction paragraph style - professional but single-page optimized
        intro_style = styles.add_style('Custom Intro', WD_STYLE_TYPE.PARAGRAPH)
        intro_font = intro_style.font
        intro_font.name = 'Calibri'
        intro_font.size = Pt(10.5)  # Slightly reduced for single page
        intro_font.color.rgb = RGBColor(0, 0, 0)
        intro_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        intro_style.paragraph_format.space_after = Pt(12)  # Reduced for single page
        intro_style.paragraph_format.space_before = Pt(8)
        intro_style.paragraph_format.line_spacing = 1.2  # Reduced but still readable
        intro_style.paragraph_format.left_indent = Inches(0.2)
        intro_style.paragraph_format.right_indent = Inches(0.2)
        
        # Body text style - compact
        body_style = styles.add_style('Custom Body', WD_STYLE_TYPE.PARAGRAPH)
        body_font = body_style.font
        body_font.name = 'Calibri'
        body_font.size = Pt(9)  # Further reduced for single page
        body_font.color.rgb = RGBColor(0, 0, 0)
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_style.paragraph_format.space_after = Pt(4)  # Reduced spacing
        body_style.paragraph_format.line_spacing = 1.15
        body_style.paragraph_format.first_line_indent = Inches(0)
        
        # Reason heading style - bold and impactful but single-page optimized
        reason_style = styles.add_style('Custom Reason', WD_STYLE_TYPE.PARAGRAPH)
        reason_font = reason_style.font
        reason_font.name = 'Calibri'
        reason_font.size = Pt(12)  # Optimized for single page
        reason_font.bold = True
        reason_font.color.rgb = RGBColor(0, 0, 0)
        reason_style.paragraph_format.space_before = Pt(6)  # Further reduced spacing between reasons
        reason_style.paragraph_format.space_after = Pt(3)  
        reason_style.paragraph_format.line_spacing = 1.15
        reason_style.paragraph_format.keep_with_next = True
        reason_style.paragraph_format.left_indent = Inches(0.15)
        
        # Reason body style - professional and readable but single-page optimized
        reason_body_style = styles.add_style('Custom Reason Body', WD_STYLE_TYPE.PARAGRAPH)
        reason_body_font = reason_body_style.font
        reason_body_font.name = 'Calibri'
        reason_body_font.size = Pt(10)  # Optimized for single page
        reason_body_font.color.rgb = RGBColor(0, 0, 0)
        reason_body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        reason_body_style.paragraph_format.space_after = Pt(5)  # Further reduced spacing after each reason
        reason_body_style.paragraph_format.line_spacing = 1.2
        reason_body_style.paragraph_format.left_indent = Inches(0.25)
        reason_body_style.paragraph_format.right_indent = Inches(0.15)
        
        # Section separator style - minimal
        separator_style = styles.add_style('Custom Separator', WD_STYLE_TYPE.PARAGRAPH)
        separator_style.paragraph_format.space_before = Pt(6)
        separator_style.paragraph_format.space_after = Pt(6)
        separator_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _parse_generated_content(self, content):
        """Parse the generated marketing content into structured components"""
        lines = content.strip().split('\n')
        
        # Find the company name from the first substantial line
        company_name = "Company"
        for line in lines[:8]:  # Look deeper in content
            line = line.strip()
            if line and not line.startswith(('Top 5', 'Here are', 'The following', '1.', '2.', '3.', '4.', '5.')):
                # Extract company name from lines that mention partnership or should partner
                if 'should partner' in line.lower() or 'partner with' in line.lower():
                    # Look for pattern like "Company Name should partner with TalentCraft"
                    match = re.search(r'([A-Za-z][A-Za-z\s&.,Inc-]+?)\s+should partner', line, re.IGNORECASE)
                    if match:
                        company_name = match.group(1).strip()
                        break
                elif len(line) > 15 and len(line) < 120:  # More flexible range
                    # Look for company names in descriptive text
                    potential_names = re.findall(r'\b([A-Z][a-zA-Z]{2,}(?:\s+[A-Z][a-zA-Z]*)*)', line)
                    if potential_names:
                        company_name = potential_names[0] if isinstance(potential_names[0], str) else potential_names[0]
                        break
        
        # Clean up company name
        company_name = re.sub(r'\s+(should|needs|faces|is)\s+.*', '', company_name, flags=re.IGNORECASE)
        company_name = company_name.strip()
        
        # Find the introduction paragraph (usually the first substantial paragraph)
        intro_paragraph = ""
        collecting_intro = False
        intro_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                if collecting_intro and intro_lines:
                    # If we've been collecting and hit empty line, we might be done
                    continue
                else:
                    continue
            
            # Skip obvious reason lines
            if re.match(r'^\d+\.', line) or line.lower().startswith(('1.', '2.', '3.', '4.', '5.')):
                break
            
            # Skip title-like lines and headers
            if (('top 5' in line.lower() and 'reasons' in line.lower()) or 
                ('talentcraft' in line.lower() and 'should' in line.lower() and 'partner' in line.lower()) or
                len(line) < 25 or
                line.isupper() or
                line.endswith(':') or
                'flexible tech' in line.lower()):
                continue
            
            # Look for substantial content that forms the introduction
            if len(line) > 30 and not collecting_intro:
                collecting_intro = True
            
            if collecting_intro:
                intro_lines.append(line)
                
                # Stop collecting if we have enough content (1-3 sentences usually)
                if len(' '.join(intro_lines)) > 200:
                    break
        
        # Join intro lines and clean up
        if intro_lines:
            intro_paragraph = ' '.join(intro_lines)
            # Clean up any residual formatting
            intro_paragraph = re.sub(r'\s+', ' ', intro_paragraph).strip()
            # Ensure it's a complete thought
            if not intro_paragraph.endswith(('.', '!', '?')):
                intro_paragraph += '.'
        
        # Extract the 5 reasons with improved parsing
        reasons = []
        current_reason = None
        current_reason_text = ""
        
        in_reasons_section = False
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if this is a numbered reason (more flexible patterns)
            reason_match = re.match(r'^(\d+)\.\s*(.+)', line)
            if reason_match:
                # Save previous reason if exists
                if current_reason is not None:
                    reasons.append({
                        'number': current_reason,
                        'title': self._extract_reason_title(current_reason_text),
                        'content': self._clean_reason_content(current_reason_text)
                    })
                
                # Start new reason
                current_reason = int(reason_match.group(1))
                current_reason_text = reason_match.group(2)
                in_reasons_section = True
                
            elif in_reasons_section and current_reason is not None:
                # Continue building current reason content
                # Stop if we hit the next numbered item
                next_reason_match = re.match(r'^(\d+)\.', line)
                if next_reason_match and int(next_reason_match.group(1)) == current_reason + 1:
                    # This is the next reason, don't include in current
                    continue
                elif not re.match(r'^\d+\.', line):  # Not a numbered line
                    current_reason_text += " " + line
        
        # Don't forget the last reason
        if current_reason is not None:
            reasons.append({
                'number': current_reason,
                'title': self._extract_reason_title(current_reason_text),
                'content': self._clean_reason_content(current_reason_text)
            })
        
        return {
            'company_name': company_name,
            'intro_paragraph': intro_paragraph.strip(),
            'reasons': reasons[:5]  # Ensure we only take first 5
        }
    
    def _extract_reason_title(self, reason_text):
        """Extract a clear title from the reason text"""
        # Look for patterns like "TITLE:" or all-caps sections first
        if ':' in reason_text:
            potential_title = reason_text.split(':')[0].strip()
            # Check if title is in all caps or has multiple uppercase words
            # And ensure it's a reasonable length for a title (not too long)
            if (len(potential_title) < 100 and len(potential_title) > 10 and
                (potential_title.isupper() or 
                 len([w for w in potential_title.split() if w.isupper()]) >= 2)):
                return potential_title
        
        # Remove markdown-style formatting
        clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', reason_text)
        clean_text = re.sub(r'\*(.*?)\*', r'\1', clean_text)
        
        # Try to extract the first sentence that ends with a period
        # Split on periods and take the first complete sentence
        sentences = clean_text.split('.')
        if sentences and len(sentences) > 1:  # Ensure there are multiple sentences
            first_sentence = sentences[0].strip()
            # Make sure it's a reasonable title length and not too long
            if 10 < len(first_sentence) <= 80:  # Reduced max length for better titles
                return first_sentence
        
        # If no good sentence found, look for the first part before a colon
        if ':' in clean_text:
            title_part = clean_text.split(':')[0].strip()
            if 10 < len(title_part) <= 80:
                return title_part
        
        # Fallback: use first significant words but keep it short
        words = clean_text.split()[:12]  # Reduced from 15
        title = ' '.join(words)
        # Ensure reasonable length
        if len(title) > 80:
            title = ' '.join(words[:8])
        
        return title
    
    def _add_header_logo_space(self):
        """Add compact header space optimized for single page"""
        # Compact header space for single page
        logo_para = self.doc.add_paragraph("[Company Logo]" + " " * 40 + "TalentCraft")
        logo_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  
        logo_para.paragraph_format.space_after = Pt(12)  # Reduced for single page
        logo_para.paragraph_format.space_before = Pt(6)  
        logo_font = logo_para.runs[0].font
        logo_font.size = Pt(8)  # Reduced for single page
        logo_font.color.rgb = RGBColor(128, 128, 128)
        logo_font.italic = True
    
    def _clean_reason_content(self, reason_text):
        """Clean and format the reason content"""
        # Remove title if it exists (before first colon)
        if ':' in reason_text:
            content = ':'.join(reason_text.split(':')[1:]).strip()
        else:
            content = reason_text.strip()
        
        # Clean up extra spaces and formatting
        content = re.sub(r'\s+', ' ', content)
        content = content.strip()
        
        return content
    
    def generate_document(self, marketing_content, company_name_override=None, job_roles=""):
        """Generate a Word document with the provided marketing content"""
        
        # Parse the content
        parsed_content = self._parse_generated_content(marketing_content)
        
        # Use override if provided
        if company_name_override:
            parsed_content['company_name'] = company_name_override
        
        # Add header space for logos (matching Template1.pdf)
        self._add_header_logo_space()
        
        # Add document identifier in top right - compact for single page
        header_para = self.doc.add_paragraph("TalentCraft Partnership Proposal", style='Custom Body')
        header_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_para.paragraph_format.space_after = Pt(8)  # Reduced for single page
        header_font = header_para.runs[0].font
        header_font.size = Pt(8)  # Reduced for single page
        header_font.color.rgb = RGBColor(100, 100, 100)
        header_font.italic = True
        
        # Create the document title (matching Template1.pdf structure)
        title = f"Top 5 Reasons {parsed_content['company_name']} Should Partner with TalentCraft"
        
        # Split title for better readability but optimized for single page
        title_lines = self._split_title_intelligently(title, parsed_content['company_name'])
        for i, line in enumerate(title_lines):
            title_para = self.doc.add_paragraph(line, style='Custom Title')
            if i == 0:  # First line
                title_para.paragraph_format.space_after = Pt(1)  # Minimal spacing
            else:  # Second line
                title_para.paragraph_format.space_before = Pt(0)
                title_para.paragraph_format.space_after = Pt(3)  # Minimal spacing
        
        # Add subtitle (matching template structure)
        if job_roles:
            subtitle = f"Flexible Tech + Leadership Talent for {job_roles}"
        else:
            subtitle = "Flexible Tech + Leadership Talent for Account Director of Sales"
        
        subtitle_para = self.doc.add_paragraph(subtitle, style='Custom Subtitle')
        
        # Add dividing line under subtitle
        self._add_dividing_line()
        
        # Add introduction paragraph with proper styling
        if parsed_content['intro_paragraph']:
            # Clean up the intro paragraph
            intro_text = self._clean_intro_paragraph(parsed_content['intro_paragraph'])
            intro_para = self.doc.add_paragraph(intro_text, style='Custom Intro')
        else:
            # Add default intro if none found
            default_intro = f"Research-Backed Proposal for {parsed_content['company_name']}: TalentCraft Partnership"
            intro_para = self.doc.add_paragraph(default_intro, style='Custom Intro')
        
        # Add the 5 reasons with improved formatting
        for i, reason in enumerate(parsed_content['reasons']):
            if i >= 5:  # Ensure we only show 5 reasons
                break
                
            # Add reason title with number (matching template style)
            reason_title = f"{reason['number']}. {reason['title']}"
            reason_para = self.doc.add_paragraph(reason_title, style='Custom Reason')
            
            # Add reason content with proper formatting
            if reason['content']:
                # Clean and format the reason content
                clean_content = self._clean_reason_content_for_doc(reason['content'])
                content_para = self.doc.add_paragraph(clean_content, style='Custom Reason Body')
                
            # No additional separators needed - spacing is handled by reason styles
        
        # No automatic closing statement needed
        
        return self.doc
    
    def _split_title_intelligently(self, title, company_name):
        """Split long titles into multiple lines intelligently (matching Template1.pdf)"""
        # Always split for better visual impact, like the template
        if "Should Partner with TalentCraft" in title:
            return [
                f"Top 5 Reasons {company_name} Should Partner with TalentCraft"
            ]
        elif len(title) > 50:  # Split longer titles
            # Find natural break point
            words = title.split()
            if "Should" in words:
                should_index = words.index("Should")
                return [
                    " ".join(words[:should_index]),
                    " ".join(words[should_index:])
                ]
            else:
                mid_point = len(words) // 2
                return [
                    " ".join(words[:mid_point]),
                    " ".join(words[mid_point:])
                ]
        else:
            return [title]
    
    def _clean_intro_paragraph(self, intro_text):
        """Clean and format introduction paragraph"""
        # Remove any remaining markdown or formatting
        intro_text = intro_text.strip()
        
        # Remove common prefixes that might appear
        prefixes_to_remove = [
            "Research-backed proposal for",
            "Introduction:",
            "Executive Summary:",
            "Overview:"
        ]
        
        for prefix in prefixes_to_remove:
            if intro_text.lower().startswith(prefix.lower()):
                intro_text = intro_text[len(prefix):].strip()
                intro_text = intro_text.lstrip(':').strip()
        
        # Ensure it ends with proper punctuation
        if intro_text and not intro_text.endswith(('.', '!', '?')):
            intro_text += "."
        
        return intro_text
    
    def _add_dividing_line(self):
        """Add a horizontal dividing line under the subtitle"""
        # Create a paragraph with a horizontal line
        line_para = self.doc.add_paragraph("", style='Custom Divider')
        
        # Add the horizontal line using underscore characters
        line_run = line_para.add_run("_" * 80)  # Create a line of underscores
        line_run.font.size = Pt(8)
        line_run.font.color.rgb = RGBColor(128, 128, 128)  # Light gray
        
        # Alternative method using border (if needed)
        # from docx.oxml.ns import qn
        # from docx.oxml import OxmlElement
        # border = OxmlElement('w:pBdr')
        # bottom = OxmlElement('w:bottom')
        # bottom.set(qn('w:val'), 'single')
        # bottom.set(qn('w:sz'), '6')
        # bottom.set(qn('w:space'), '1')
        # bottom.set(qn('w:color'), 'auto')
        # border.append(bottom)
        # line_para._element.get_or_add_pPr().append(border)
    
    def _clean_reason_content_for_doc(self, content):
        """Clean reason content for document formatting"""
        # Remove any markdown formatting
        content = content.strip()
        
        # Remove "Why it matters:" sections that might be redundant
        content = re.sub(r'\s*Why it matters:\s*', ' ', content, flags=re.IGNORECASE)
        
        # Clean up multiple spaces and ensure proper punctuation
        content = re.sub(r'\s+', ' ', content)
        
        if content and not content.endswith(('.', '!', '?')):
            content += "."
        
        return content
    
    def _add_closing_statement(self, company_name, job_roles):
        """Add compact closing statement optimized for single page"""
        # Minimal spacing before closing for single page
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)  # Reduced for single page
        
        # Compact professional closing statement
        closing_text = f"TalentCraft's proven expertise in {job_roles} recruitment positions {company_name} for immediate competitive advantage through strategic talent acquisition."
        
        closing_para = self.doc.add_paragraph(closing_text, style='Custom Body')
        closing_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        closing_para.paragraph_format.space_before = Pt(8)  # Reduced for single page
        closing_para.paragraph_format.space_after = Pt(6)  
        closing_para.paragraph_format.left_indent = Inches(0.2)  
        closing_para.paragraph_format.right_indent = Inches(0.2)  
        
        # Style the closing statement
        closing_font = closing_para.runs[0].font
        closing_font.size = Pt(9.5)  # Reduced for single page
        closing_font.color.rgb = RGBColor(64, 64, 64)
        closing_font.italic = True
    

    
    def save_document(self, file_path):
        """Save the document to the specified path"""
        if self.doc:
            self.doc.save(file_path)
            return True
        return False
    
    def create_document_from_content(self, marketing_content, company_name=None, job_roles="", output_path=None):
        """Complete workflow: parse content, generate document, and optionally save"""
        
        # Generate the document
        doc = self.generate_document(marketing_content, company_name, job_roles)
        
        # Save if path provided
        if output_path:
            self.save_document(output_path)
            return output_path
        
        return doc

def generate_marketing_document(marketing_content, company_name=None, job_roles="", output_filename=None):
    """
    Convenience function to generate a marketing document
    
    Args:
        marketing_content (str): The generated marketing content from the app
        company_name (str, optional): Override company name if needed
        job_roles (str, optional): Job roles for subtitle
        output_filename (str, optional): Filename to save document
    
    Returns:
        str: Path to saved document or None if not saved
    """
    
    generator = TalentCraftDocumentGenerator()
    
    # Generate safe filename if not provided
    if output_filename is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_company = re.sub(r'[^\w\s-]', '', company_name or "Company").strip()
        safe_company = re.sub(r'[-\s]+', '_', safe_company)
        output_filename = f"TalentCraft_Proposal_{safe_company}_{timestamp}.docx"
    
    # Ensure .docx extension
    if not output_filename.endswith('.docx'):
        output_filename += '.docx'
    
    # Create full path
    output_path = os.path.join(os.getcwd(), output_filename)
    
    # Generate and save document
    generator.create_document_from_content(
        marketing_content=marketing_content,
        company_name=company_name,
        job_roles=job_roles,
        output_path=output_path
    )
    
    return output_path

if __name__ == "__main__":
    # Test with sample content
    sample_content = """
    TechCorp faces significant challenges in securing top-tier talent for Software Engineering and Data Science positions in today's competitive market.

    1. SPECIALIZED EXPERTISE: We understand the unique requirements for Software Engineering and Data Science roles and have a proven track record of successful placements in the tech industry.

    2. EXTENSIVE NETWORK: Our talent network includes top-tier professionals specifically in your target roles with experience in Python, machine learning, and cloud technologies.

    3. STREAMLINED PROCESS: We handle the entire recruitment process, saving your team valuable time and resources while ensuring quality candidates.

    4. QUALITY GUARANTEE: We ensure each candidate meets your exact specifications before presentation, with thorough technical vetting.

    5. ONGOING SUPPORT: Our partnership doesn't end with placement - we provide ongoing support to ensure long-term success and team integration.
    """
    
    output_path = generate_marketing_document(
        marketing_content=sample_content,
        company_name="TechCorp",
        job_roles="Software Engineering & Data Science"
    )
    
    print(f"Test document generated: {output_path}")
